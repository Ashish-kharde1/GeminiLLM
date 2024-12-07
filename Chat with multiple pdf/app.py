import streamlit as st
from PyPDF2 import PdfReader
from langchain.text_splitter import RecursiveCharacterTextSplitter
import os
from langchain_google_genai import GoogleGenerativeAIEmbeddings
import google.generativeai as genai
from langchain.vectorstores import FAISS
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.chains.question_answering import load_qa_chain
from langchain.prompts import PromptTemplate
from dotenv import load_dotenv
import json
import pandas as pd
import docx

load_dotenv()
os.getenv("GOOGLE_API_KEY")
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))






# def extract_data_from_pdf(pdf_docs): 
#     text=""
#     for pdf in pdf_docs:
#         pdf_reader= PdfReader(pdf)
#         for page in pdf_reader.pages:
#             text+= page.extract_text()
#     return  text

# def extract_data_from_json(json_docs):
#     text = ""
#     for json_file in json_docs:
#         json_data = json.load(json_file)  # Load the JSON data from the file
#         # Convert the entire JSON data into a string
#         text += json.dumps(json_data, indent=4)  # This adds a formatted string representation of the JSON
#     return text

# def extract_data_from_txt(txt_docs):
#     text = ""
#     for txt_file in txt_docs:
#         text += txt_file.read().decode("utf-8")  # Decode as UTF-8
#     return text

# def extract_data_from_csv(csv_docs):
#     text = ""
#     for csv_file in csv_docs:
#         df = pd.read_csv(csv_file)  # Load CSV data into a DataFrame
#         text += df.to_string(index=False)  # Convert DataFrame to string
#     return text

# def extract_data_from_xlsx(xlsx_docs):
#     text = ""
#     for xlsx_file in xlsx_docs:
#         df = pd.read_excel(xlsx_file)  # Load Excel data into a DataFrame
#         text += df.to_string(index=False)  # Convert DataFrame to string
#     return text

# def extract_data_from_docx(docx_docs):
#     text = ""
#     for docx_file in docx_docs:
#         doc = docx.Document(docx_file)  # Load DOCX file
#         for para in doc.paragraphs:
#             text += para.text  # Extract text from paragraphs
#     return text


# def extract_all_data(uploaded_files):
#     all_data = ""
    
#     # Separate files by type
#     pdf_files = [file for file in uploaded_files if file.name.endswith('.pdf')]
#     json_files = [file for file in uploaded_files if file.name.endswith('.json')]
#     txt_files = [file for file in uploaded_files if file.name.endswith('.txt')]
#     csv_files = [file for file in uploaded_files if file.name.endswith('.csv')]
#     xlsx_files = [file for file in uploaded_files if file.name.endswith('.xlsx')]
#     docx_files = [file for file in uploaded_files if file.name.endswith('.docx')]


#     # Extract data for each file type
#     if pdf_files:
#         all_data += extract_data_from_pdf(pdf_files)
#     if json_files:
#         all_data += extract_data_from_json(json_files)
#     if txt_files:
#         all_data += extract_data_from_txt(txt_files)
#     if csv_files:
#         all_data += extract_data_from_csv(csv_files)
#     if xlsx_files:
#         all_data += extract_data_from_xlsx(xlsx_files)
#     if docx_files:
#         all_data += extract_data_from_docx(docx_files)


#     return all_data


def extract_all_data(uploaded_files):
    all_data = ""

    # Iterate through each uploaded file
    for file in uploaded_files:
        file_type = file.name.split('.')[-1].lower()  # Get the file extension

        # Process PDF files
        if file_type == 'pdf':
            pdf_reader = PdfReader(file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text()
            all_data += f"--- Text extracted from {file.name} (PDF) ---\n{text}\n\n"

        # Process JSON files
        elif file_type == 'json':
            json_data = json.load(file)  # Load JSON data
            text = json.dumps(json_data, indent=4)  # Convert JSON to a formatted string
            all_data += f"--- Text extracted from {file.name} (JSON) ---\n{text}\n\n"

        # Process TXT files
        elif file_type == 'txt':
            text = file.read().decode("utf-8")  # Read text from file
            all_data += f"--- Text extracted from {file.name} (TXT) ---\n{text}\n\n"

        # Process CSV files
        elif file_type == 'csv':
            df = pd.read_csv(file)  # Load CSV into a DataFrame
            text = df.to_string(index=False)  # Convert DataFrame to string
            all_data += f"--- Data extracted from {file.name} (CSV) ---\n{text}\n\n"

        # Process XLSX files
        elif file_type == 'xlsx':
            df = pd.read_excel(file)  # Load Excel into a DataFrame
            text = df.to_string(index=False)  # Convert DataFrame to string
            all_data += f"--- Data extracted from {file.name} (XLSX) ---\n{text}\n\n"

        # Process DOCX files
        elif file_type == 'docx':
            doc = docx.Document(file)  # Load DOCX file
            text = ""
            for para in doc.paragraphs:
                text += para.text  # Extract text from paragraphs
            all_data += f"--- Text extracted from {file.name} (DOCX) ---\n{text}\n\n"

        else:
            all_data += f"--- Unsupported file type: {file.name} ---\n"

    return all_data


def get_text_chunks(text):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=10000, chunk_overlap=1000)
    chunks = text_splitter.split_text(text)
    return chunks


def get_vector_store(text_chunks):
    embeddings = GoogleGenerativeAIEmbeddings(model = "models/embedding-001")
    vector_store = FAISS.from_texts(text_chunks, embedding=embeddings)
    vector_store.save_local("faiss_index")


def get_conversational_chain():

    prompt_template = """
    Answer the question as detailed as possible from the provided context, make sure to provide all the details, if the answer is not in
    provided context just say, "answer is not available in the context", don't provide the wrong answer\n\n
    Context:\n {context}?\n
    Question: \n{question}\n

    Answer:
    """

    model = ChatGoogleGenerativeAI(model="gemini-1.5-flash",
                             temperature=0.3)

    prompt = PromptTemplate(template = prompt_template, input_variables = ["context", "question"])
    chain = load_qa_chain(model, chain_type="stuff", prompt=prompt)

    return chain



def user_input(user_question):
    embeddings = GoogleGenerativeAIEmbeddings(model = "models/embedding-001")
    
    new_db = FAISS.load_local("faiss_index", embeddings, allow_dangerous_deserialization=True)
    docs = new_db.similarity_search(user_question)

    chain = get_conversational_chain()

    
    response = chain(
        {"input_documents":docs, "question": user_question}
        , return_only_outputs=True)

    print(response)
    st.write("Reply: ", response["output_text"])




def main():
    st.set_page_config("Chat with Documents")
    st.header("Chat with Documents using Gemini💁")

    user_question = st.text_input("Ask a Question from the Document Files")

    if user_question:
        user_input(user_question)

    with st.sidebar:
        st.title("Menu:")
        # pdf_docs = st.file_uploader("Upload your PDF Files and Click on the Submit & Process Button", accept_multiple_files=True)
        pdf_docs = st.file_uploader("Upload your files (PDF, TXT, DOCX, etc.) and Click on the Submit & Process Button", accept_multiple_files=True, type=["pdf", "txt", "docx", "csv", "xlsx", "json"]  # Add the file types you want to accept
)
        if st.button("Submit & Process"):
            with st.spinner("Processing..."):
                raw_text = extract_all_data(pdf_docs)
                text_chunks = get_text_chunks(raw_text)
                get_vector_store(text_chunks)
                st.success("Done")



if __name__ == "__main__":
    main()